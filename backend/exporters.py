from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import LETTER
import re

def add_bold(paragraph,text):

    parts=re.split(r"(\*\*.*?\*\*)",text)

    for p in parts:

        if p.startswith("**") and p.endswith("**"):
            run=paragraph.add_run(p[2:-2])
            run.bold=True
        else:
            paragraph.add_run(p)


def export_docx(cv,path):

    doc=Document()

    h=cv["header"]

    doc.add_heading(h["name"],0)

    doc.add_paragraph(h["title"])

    doc.add_paragraph(f"{h['location']} | {h['email']} | {h['phone']}")

    doc.add_heading("Summary",1)

    p=doc.add_paragraph()

    add_bold(p,cv["summary"])

    doc.add_heading("Experience",1)

    for e in cv["experience"]:

        doc.add_paragraph(f"{e['role']} — {e['company']}")

        for b in e["bullets"]:

            p=doc.add_paragraph(style="List Bullet")

            add_bold(p,b)

    doc.add_heading("Projects",1)

    for p in cv["projects"]:

        doc.add_paragraph(p["name"])

        for b in p["bullets"]:

            doc.add_paragraph(b,style="List Bullet")

    doc.save(path)



def export_pdf(cv,path):

    c=canvas.Canvas(path,pagesize=LETTER)

    y=750

    pages=1

    def write(t):

        nonlocal y,pages

        if y<80:

            if pages>=2:
                return False

            c.showPage()

            pages+=1

            y=750

        c.drawString(50,y,t)

        y-=15

        return True


    write(cv["header"]["name"])

    write(cv["summary"])

    for e in cv["experience"]:

        write(e["role"])

        for b in e["bullets"]:

            write(b)

    c.save()