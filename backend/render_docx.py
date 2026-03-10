# backend/render_docx.py
from __future__ import annotations

from pathlib import Path
from typing import Any, Dict, List
import math
import re

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph

_BOLD_RE = re.compile(r"\*\*(.+?)\*\*")


def _strip_md_bold_markers(text: str) -> str:
    if not text:
        return ""
    return re.sub(r"\*\*(.+?)\*\*", r"\1", str(text)).strip()


# ============================================================
# Markdown **bold** -> Word bold runs
# ============================================================
def _add_md_bold_runs(par: Paragraph, text: str) -> None:
    s = text or ""
    pos = 0
    for m in _BOLD_RE.finditer(s):
        if m.start() > pos:
            par.add_run(s[pos:m.start()])
        r = par.add_run(m.group(1))
        r.bold = True
        pos = m.end()
    if pos < len(s):
        par.add_run(s[pos:])


# ============================================================
# Public API
# ============================================================
def render_docx(out_path: str | Path, cv: Dict[str, Any]) -> Path:
    out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    _set_page_margins(doc, top=0.55, bottom=0.55, left=0.7, right=0.7)
    _set_default_font(doc, name="Calibri", size_pt=10)

    header = _as_dict(cv.get("header"))
    _render_header(doc, header, cv)

    # 1) Professional Summary
    summary = (cv.get("summary") or "").strip()
    if summary:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run("PROFESSIONAL SUMMARY")
        r.bold = True
        r.font.size = Pt(11)

        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.line_spacing = 1.15
        _add_md_bold_runs(p, summary)

    # 2) Areas of Expertise
    areas = _as_list(cv.get("areas_of_expertise"))
    if areas:
        _section_title(doc, "AREAS OF EXPERTISE")
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.line_spacing = 1.1
        clean = [str(x).strip() for x in areas if str(x).strip()]
        p.add_run(" • ".join(clean))

    # 3) Career Highlights
    highlights = _as_list(cv.get("career_highlights"))
    if highlights:
        _section_title(doc, "CAREER HIGHLIGHTS")
        _render_two_column_highlights(doc, highlights)

    # 4) Skills
    skills = _as_dict(cv.get("skills"))
    if skills:
        _section_title(doc, "SKILLS")
        for cat, items in skills.items():
            line = ""
            if isinstance(items, list):
                clean = [str(x).strip() for x in items if str(x).strip()]
                line = ", ".join(clean)
            else:
                line = str(items).strip()

            if not line:
                continue

            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(5)
            p.paragraph_format.line_spacing = 1.1
            r1 = p.add_run(f"{cat}: ")
            r1.bold = True
            p.add_run(line)

    # 5) Work Experience
    experience = _as_list_of_dicts(cv.get("experience"))
    if experience:
        _section_title(doc, "WORK EXPERIENCE")
        for ex in experience:
            exp_dates = str(ex.get("dates") or "").strip()
            if not exp_dates:
                s = str(ex.get("start") or "").strip()
                e = str(ex.get("end") or "").strip()
                exp_dates = f"{s} – {e}".strip(" –")

            _render_role_company_block(
                doc,
                role=str(ex.get("role") or "").strip(),
                company=str(ex.get("company") or "").strip(),
                dates=exp_dates,
                bullets=_as_list(ex.get("bullets")),
            )

    # 6) Projects
    projects = _as_list_of_dicts(cv.get("projects"))
    if projects:
        _section_title(doc, "PROJECTS")
        for pr in projects:
            _render_project_item(doc, pr)

    # 7) Education
    education = _as_list_of_dicts(cv.get("education"))
    if education:
        _section_title(doc, "EDUCATION")
        for ed in education:
            _render_education_item(doc, ed)

    # 8) Certifications
    certs_any = cv.get("certifications")
    certs_list = _as_list(certs_any)
    if certs_list:
        _section_title(doc, "CERTIFICATIONS")
        _render_certifications(doc, certs_list)

    # 9) Internships
    internships = _as_list_of_dicts(cv.get("internships"))
    if internships:
        _section_title(doc, "INTERNSHIPS")
        for it in internships:
            exp_dates = str(ex.get("dates") or "").strip()
            if not exp_dates:
                s = str(ex.get("start") or "").strip()
                e = str(ex.get("end") or "").strip()
                exp_dates = f"{s} – {e}".strip(" –")

            _render_role_company_block(
                doc,
                role=str(ex.get("role") or "").strip(),
                company=str(ex.get("company") or "").strip(),
                dates=exp_dates,
                bullets=_as_list(ex.get("bullets")),
            )

    doc.save(str(out_path))
    return out_path


# ============================================================
# Header
# ============================================================
def _render_header(doc: Document, header: Dict[str, Any], cv: Dict[str, Any]) -> None:
    name_raw = (header.get("name") or "").strip()
    name = _strip_md_bold_markers(name_raw) or "CANDIDATE NAME"

    title = str(header.get("title") or cv.get("title") or "").strip()

    email = (header.get("email") or "").strip()
    phone = (header.get("phone") or "").strip()
    linkedin = (header.get("linkedin") or "").strip()
    github = (header.get("github") or "").strip()
    address = (header.get("address") or header.get("location") or "").strip()

    # Name
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(name.upper())
    r.bold = True
    r.font.size = Pt(17)

    # Professional title
    if title:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(8)
        r = p.add_run(title)
        r.bold = True
        r.font.size = Pt(11)
    

    # Address | Email | Phone
    line1_parts = [x for x in [address, email, phone] if x]
    if line1_parts:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(" | ".join(line1_parts))
        r.font.size = Pt(10)

    # GitHub | LinkedIn
    line2_parts = [x for x in [github, linkedin] if x]
    if line2_parts:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(" | ".join(line2_parts))
        r.font.size = Pt(10)

    _horizontal_rule(doc, color="666666", space_before=0, space_after=14)


# ============================================================
# Sections
# ============================================================
def _section_title(doc: Document, title: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(11)

    _horizontal_rule(doc, color="777777", space_before=0, space_after=5)


# ============================================================
# Education / Experience / Projects / Internships blocks
# ============================================================
def _render_education_item(doc: Document, ed: Dict[str, Any]) -> None:
    degree = (ed.get("degree") or "").strip()
    university = (ed.get("university") or "").strip()
    city = (ed.get("city") or "").strip()
    country = (ed.get("country") or "").strip()
    dates = (ed.get("dates") or "").strip()

    line = (ed.get("line") or "").strip()
    if not (degree or university) and line:
        _two_col_line(doc, left_text=line, right_text="", left_bold=True)
        cw = (ed.get("coursework") or "").strip()
        honors = (ed.get("honors") or "").strip()
        if cw:
            _bullet(doc, f"Relevant Coursework: {cw}")
        if honors:
            _bullet(doc, f"Honors: {honors}")
        _small_gap(doc, after=5)
        return

    if degree:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(degree)
        r.bold = True

    loc = ", ".join([x for x in [university, city, country] if x]).strip()
    _two_col_line(doc, left_text=loc, right_text=dates, left_bold=False, italic_left=True)

    cw = (ed.get("coursework") or "").strip()
    honors = (ed.get("honors") or "").strip()
    if cw:
        _bullet(doc, f"Relevant Coursework: {cw}")
    if honors:
        _bullet(doc, f"Honors: {honors}")

    _small_gap(doc, after=5)


def _render_role_company_block(
    doc: Document,
    *,
    role: str,
    company: str,
    dates: str,
    bullets: List[Any],
) -> None:
    role = role or "Role"
    company = company or ""
    dates = dates or ""

    _two_col_line(doc, left_text=role, right_text=dates, left_bold=True)

    if company:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run(company)
        r.italic = True

    for b in bullets or []:
        s = str(b).strip()
        if s:
            _bullet_rich(doc, s)

    _small_gap(doc, after=5)


def _render_project_item(doc: Document, pr: Dict[str, Any]) -> None:
    name = (pr.get("name") or "").strip() or "Project"
    tech = (pr.get("tech") or "").strip()
    dates = (pr.get("dates") or "").strip()
    bullets = _as_list(pr.get("bullets"))

    _two_col_line(doc, left_text=name, right_text=dates, left_bold=True)

    if tech:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        r1 = p.add_run("Technologies: ")
        r1.bold = True
        p.add_run(tech)

    for b in bullets or []:
        s = str(b).strip()
        if s:
            _bullet_rich(doc, s)

    _small_gap(doc, after=5)


# ============================================================
# Career Highlights - Two Columns
# ============================================================
def _render_two_column_highlights(doc: Document, highlights: List[Any]) -> None:
    clean = [str(x).strip() for x in highlights if str(x).strip()]
    if not clean:
        return

    half = math.ceil(len(clean) / 2)
    left_items = clean[:half]
    right_items = clean[half:]

    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.columns[0].width = Inches(3.25)
    table.columns[1].width = Inches(3.25)

    _remove_table_borders(table)

    left_cell = table.cell(0, 0)
    right_cell = table.cell(0, 1)
    left_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    right_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

    # left
    for idx, item in enumerate(left_items):
        p = left_cell.paragraphs[0] if idx == 0 else left_cell.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.18)
        p.paragraph_format.first_line_indent = Inches(-0.12)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.08
        p.add_run("• ")
        _add_md_bold_runs(p, item)

    # right
    for idx, item in enumerate(right_items):
        p = right_cell.paragraphs[0] if idx == 0 else right_cell.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.18)
        p.paragraph_format.first_line_indent = Inches(-0.12)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.08
        p.add_run("• ")
        _add_md_bold_runs(p, item)

    _small_gap(doc, after=5)


# ============================================================
# Certifications
# ============================================================
def _render_certifications(doc: Document, certs_list: List[Any]) -> None:
    for c in certs_list or []:
        if isinstance(c, dict):
            title = (c.get("title") or c.get("name") or "").strip()
            issuer = (c.get("issuer") or "").strip()

            dates = (c.get("dates") or "").strip()
            if not dates:
                s = (c.get("start") or "").strip()
                e = (c.get("end") or "").strip()
                dates = f"{s} – {e}".strip(" –")

            left = title or "Certification"
            if issuer:
                left = f"{left} — {issuer}"

            _two_col_bullet_line(doc, left_text=left, right_text=dates, left_bold=False)
        else:
            s = str(c).strip()
            if s:
                _bullet(doc, s)

    _small_gap(doc, after=5)


# ============================================================
# Layout primitives
# ============================================================
def _two_col_line(
    doc: Document,
    *,
    left_text: str,
    right_text: str,
    left_bold: bool = False,
    italic_left: bool = False,
) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    left_w = Inches(5.6)
    right_w = Inches(1.1)
    table.columns[0].width = left_w
    table.columns[1].width = right_w

    _remove_table_borders(table)

    cell_l = table.cell(0, 0)
    cell_r = table.cell(0, 1)
    cell_l.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    cell_r.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

    p1 = cell_l.paragraphs[0]
    p1.paragraph_format.space_after = Pt(0)
    p1.paragraph_format.line_spacing = 1.0
    r1 = p1.add_run((left_text or "").strip())
    r1.bold = bool(left_bold)
    r1.italic = bool(italic_left)
    r1.font.size = Pt(10.5 if left_bold else 10)

    p2 = cell_r.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.0
    p2.add_run((right_text or "").strip())


def _two_col_bullet_line(
    doc: Document,
    *,
    left_text: str,
    right_text: str,
    left_bold: bool = False,
) -> None:
    left_text = (left_text or "").strip()
    right_text = (right_text or "").strip()
    if not left_text and not right_text:
        return

    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    left_w = Inches(5.6)
    right_w = Inches(1.1)
    table.columns[0].width = left_w
    table.columns[1].width = right_w

    _remove_table_borders(table)

    cell_l = table.cell(0, 0)
    cell_r = table.cell(0, 1)

    p1 = cell_l.paragraphs[0]
    p1.paragraph_format.space_after = Pt(2)
    p1.paragraph_format.line_spacing = 1.05
    p1.add_run("• ")
    rtxt = p1.add_run(left_text)
    rtxt.bold = bool(left_bold)

    p2 = cell_r.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p2.paragraph_format.space_after = Pt(2)
    p2.paragraph_format.line_spacing = 1.05
    p2.add_run(right_text)


def _bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.22)
    p.paragraph_format.first_line_indent = Inches(-0.14)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.08
    p.add_run("• ")
    p.add_run(text)


def _bullet_rich(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.22)
    p.paragraph_format.first_line_indent = Inches(-0.14)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.08
    p.add_run("• ")
    _add_md_bold_runs(p, text)


def _small_gap(doc: Document, after: int = 4) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)


def _horizontal_rule(doc: Document, color: str = "999999", space_before: int = 4, space_after: int = 4) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


# ============================================================
# DOCX formatting helpers
# ============================================================
def _set_page_margins(doc: Document, *, top: float, bottom: float, left: float, right: float) -> None:
    sec = doc.sections[0]
    sec.top_margin = Inches(top)
    sec.bottom_margin = Inches(bottom)
    sec.left_margin = Inches(left)
    sec.right_margin = Inches(right)


def _set_default_font(doc: Document, *, name: str, size_pt: int) -> None:
    style = doc.styles["Normal"]
    style.font.name = name
    style.font.size = Pt(size_pt)


def _remove_table_borders(table) -> None:
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = tblPr.first_child_found_in("w:tblBorders")
    if tblBorders is None:
        tblBorders = OxmlElement("w:tblBorders")
        tblPr.append(tblBorders)

    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = tblBorders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            tblBorders.append(el)
        el.set(qn("w:val"), "nil")


# ============================================================
# Coercion helpers
# ============================================================
def _as_dict(x: Any) -> Dict[str, Any]:
    return x if isinstance(x, dict) else {}


def _as_list(x: Any) -> List[Any]:
    if x is None:
        return []
    if isinstance(x, list):
        return x
    return [x]


def _as_list_of_dicts(x: Any) -> List[Dict[str, Any]]:
    if x is None:
        return []
    if isinstance(x, list):
        return [i for i in x if isinstance(i, dict)]
    if isinstance(x, dict):
        return [x]
    return []