from __future__ import annotations

from pathlib import Path
from typing import Dict, Any
import re

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

_BOLD_RE = re.compile(r"\*\*(.+?)\*\*")


def _add_md_bold_runs(par, text: str) -> None:
    s = text or ""
    pos = 0
    for m in _BOLD_RE.finditer(s):
        if m.start() > pos:
            par.add_run(s[pos:m.start()])
        r = par.add_run(m.group(1))
        r.bold = True
        pos = m.end()
    if pos < len(s):
        par.add_run(s[pos:])


def _set_page_margins(doc: Document, *, top: float, bottom: float, left: float, right: float) -> None:
    sec = doc.sections[0]
    sec.top_margin = Inches(top)
    sec.bottom_margin = Inches(bottom)
    sec.left_margin = Inches(left)
    sec.right_margin = Inches(right)


def _set_default_font(doc: Document, *, name: str = "Calibri", size_pt: int = 11) -> None:
    style = doc.styles["Normal"]
    style.font.name = name
    style.font.size = Pt(size_pt)


def render_cover_docx(
    out_path: str | Path,
    cv: Dict[str, Any],
    cover_letter: str,
    company: str = "",
    job_title: str = "",
) -> Path:
    out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    _set_page_margins(doc, top=0.7, bottom=0.7, left=0.8, right=0.8)
    _set_default_font(doc, name="Calibri", size_pt=11)

    header = cv.get("header") if isinstance(cv.get("header"), dict) else {}
    name = (header.get("name") or "").strip()
    email = (header.get("email") or "").strip()
    phone = (header.get("phone") or "").strip()
    linkedin = (header.get("linkedin") or "").strip()
    github = (header.get("github") or "").strip()
    address = (header.get("address") or header.get("location") or "").strip()

    if name:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(name)
        r.bold = True
        r.font.size = Pt(14)

    for line in [
        f"Phone: {phone}" if phone else "",
        f"Email: {email}" if email else "",
        f"LinkedIn: {linkedin}" if linkedin else "",
        f"GitHub: {github}" if github else "",
        f"Address: {address}" if address else "",
    ]:
        if line:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            p.add_run(line)

    if job_title or company:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        text = " | ".join([x for x in [job_title.strip(), company.strip()] if x])
        r = p.add_run(text)
        r.italic = True

    paragraphs = [x.strip() for x in (cover_letter or "").split("\n\n") if x.strip()]
    for para_text in paragraphs:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        _add_md_bold_runs(p, para_text)

    doc.save(str(out_path))
    return out_path